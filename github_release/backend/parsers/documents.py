from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from threading import Lock
from typing import Any
from zipfile import BadZipFile, ZipFile

from docx import Document
from PIL import Image, UnidentifiedImageError
from pypdf import PdfReader

from backend.config import MAX_UPLOAD_BYTES
from backend.errors import DocumentError


MAX_ARCHIVE_UNCOMPRESSED_BYTES = 80 * 1024 * 1024
ALLOWED = {
    "resume": {".pdf", ".docx", ".txt", ".md", ".markdown", ".png", ".jpg", ".jpeg", ".webp"},
    "job": {".pdf", ".docx", ".txt", ".md", ".markdown", ".png", ".jpg", ".jpeg", ".webp"},
}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
OCR_FAILURE_MESSAGE = "图片文字识别失败，请上传更清晰的图片或重新上传。"
_OCR_ENGINE = None
_OCR_LOCK = Lock()


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    method: str
    warnings: tuple[str, ...] = ()
    source_type: str = "text"
    ocr_used: bool = False
    ocr_text: str = ""
    ocr_metadata: dict[str, Any] = field(default_factory=dict)


def _normalise(text: str) -> str:
    cleaned = "\n".join(line.strip() for line in text.replace("\x00", "").splitlines() if line.strip())
    if not cleaned:
        raise DocumentError("没有读取到有效文字，请检查内容后重试。")
    return cleaned


def _decode(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentError("文本编码无法识别。请另存为 UTF-8 后重试。")


def _check_archive(raw: bytes) -> None:
    try:
        with ZipFile(BytesIO(raw)) as archive:
            size = sum(item.file_size for item in archive.infolist())
    except BadZipFile as exc:
        raise DocumentError("DOCX 文件已损坏或格式不正确。") from exc
    if size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
        raise DocumentError("文档解压后体积过大，请移除无关图片后重试。")


def _ocr_engine():
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        with _OCR_LOCK:
            if _OCR_ENGINE is None:
                try:
                    from rapidocr import RapidOCR
                    _OCR_ENGINE = RapidOCR()
                except Exception as exc:
                    raise DocumentError("本地 OCR 组件未正确安装，请重新安装项目依赖。") from exc
    return _OCR_ENGINE


def _ocr_image(image: Image.Image) -> tuple[str, dict[str, Any]]:
    try:
        import numpy as np
        result = _ocr_engine()(np.asarray(image.convert("RGB")))
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError(OCR_FAILURE_MESSAGE) from exc
    texts = getattr(result, "txts", None)
    scores = getattr(result, "scores", None)
    if texts is None and isinstance(result, (tuple, list)) and result:
        rows = result[0] if isinstance(result[0], list) else result
        texts = [row[1] for row in rows if isinstance(row, (tuple, list)) and len(row) >= 2]
        scores = [row[2] for row in rows if isinstance(row, (tuple, list)) and len(row) >= 3]
    values = [str(item).strip() for item in (texts or []) if str(item).strip()]
    text = "\n".join(values)
    if not re.search(r"[\w\u3400-\u9fff]", text):
        raise DocumentError(OCR_FAILURE_MESSAGE)
    numeric_scores = [float(item) for item in (scores or []) if isinstance(item, (int, float))]
    return text, {
        "engine": "RapidOCR",
        "line_count": len(values),
        "average_score": round(sum(numeric_scores) / len(numeric_scores), 4) if numeric_scores else None,
    }


def _parse_image(raw: bytes) -> tuple[str, dict[str, Any]]:
    try:
        with Image.open(BytesIO(raw)) as image:
            return _ocr_image(image)
    except DocumentError:
        raise
    except (UnidentifiedImageError, OSError) as exc:
        raise DocumentError("图片无法读取，请确认文件没有损坏。") from exc


def _valid_embedded_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    return len(compact) >= 30 and bool(re.search(r"[\w\u3400-\u9fff]", compact))


def _parse_pdf(raw: bytes) -> tuple[str, bool, str, dict[str, Any], tuple[str, ...]]:
    try:
        reader = PdfReader(BytesIO(raw))
        if reader.is_encrypted:
            raise DocumentError("PDF 已加密，请解除密码后重新上传。")
        embedded = [(page.extract_text() or "") for page in reader.pages]
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError("PDF 无法读取，可能已损坏。") from exc

    missing = [index for index, text in enumerate(embedded) if not _valid_embedded_text(text)]
    if not missing:
        return "\n".join(embedded), False, "", {"pages": len(embedded), "ocr_pages": []}, ()
    try:
        import pymupdf
        document = pymupdf.open(stream=raw, filetype="pdf")
    except Exception as exc:
        raise DocumentError("扫描型 PDF 无法渲染，请确认文件未损坏。") from exc
    page_texts = list(embedded)
    ocr_parts: list[str] = []
    scores: list[float] = []
    for index in missing:
        page = document[index]
        pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2.2, 2.2), alpha=False)
        with Image.open(BytesIO(pixmap.tobytes("png"))) as image:
            text, metadata = _ocr_image(image)
        page_texts[index] = text
        ocr_parts.append(text)
        if metadata.get("average_score") is not None:
            scores.append(metadata["average_score"])
    document.close()
    combined = "\n".join(page_texts)
    if not re.search(r"[\w\u3400-\u9fff]", combined):
        raise DocumentError(OCR_FAILURE_MESSAGE)
    metadata = {
        "engine": "RapidOCR", "pages": len(page_texts), "ocr_pages": [index + 1 for index in missing],
        "average_score": round(sum(scores) / len(scores), 4) if scores else None,
    }
    return combined, True, "\n".join(ocr_parts), metadata, (f"已对 {len(missing)} 页扫描内容执行本地 OCR。",)


def _parse_docx(raw: bytes) -> str:
    _check_archive(raw)
    try:
        document = Document(BytesIO(raw))
    except Exception as exc:
        raise DocumentError("DOCX 无法读取，请确认文件未损坏。") from exc
    parts = [item.text for item in document.paragraphs if item.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def parse_document(*, filename: str, raw: bytes, kind: str) -> ParsedDocument:
    if kind not in ALLOWED:
        raise DocumentError("未知文档类型。")
    if not raw:
        raise DocumentError("上传文件为空。")
    if len(raw) > MAX_UPLOAD_BYTES:
        raise DocumentError("文件超过 12 MB，请压缩后重试。")
    suffix = Path(filename).suffix.casefold()
    if suffix not in ALLOWED[kind]:
        raise DocumentError(f"该区域仅支持 {'、'.join(sorted(ALLOWED[kind]))}。")
    if suffix in IMAGE_SUFFIXES:
        text, metadata = _parse_image(raw)
        return ParsedDocument(
            text=_normalise(text), method="local_ocr", source_type="image", ocr_used=True,
            ocr_text=text, ocr_metadata=metadata,
        )
    if suffix == ".pdf":
        text, ocr_used, ocr_text, metadata, warnings = _parse_pdf(raw)
        return ParsedDocument(
            text=_normalise(text), method="pdf_mixed_ocr" if ocr_used else "pdf_embedded_text",
            source_type="file", ocr_used=ocr_used, ocr_text=ocr_text,
            ocr_metadata=metadata, warnings=warnings,
        )
    if suffix == ".docx":
        text, method = _parse_docx(raw), "docx_text_and_tables"
    else:
        text, method = _decode(raw), "plain_text"
    return ParsedDocument(text=_normalise(text), method=method, source_type="file")


def parse_text(text: str, *, kind: str) -> ParsedDocument:
    if kind not in ALLOWED:
        raise DocumentError("未知文本类型。")
    return ParsedDocument(text=_normalise(text), method="user_text", source_type="text")
